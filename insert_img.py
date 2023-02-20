#########################################################################
#Author: nilecui
#Date: 2023-02-20 10:54:59
#LastEditors: nilecui
#LastEditTime: 2023-02-20 13:21:39
#FilePath: /kafka_wk/docs_table_wk/insert_img.py
#Description: 
#Details do not determine success or failure!
#Copyright (c) 2023 by nilecui, All Rights Reserved. 
#########################################################################
import errno
import os
from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #导入的关于word操作的相关模块
import datetime #获取当前日期
from glob import glob

class SaveImgInTable:
    def __init__(self, src_dir="tables", save_dir="results", image_dir="images",img_flag="<<img>>"):
        """获取所有doc文档列表，批量处理，存储到save_dir目录
           注意：标的必须是有图片的，和文档对应，名字相同, 
                 图片所有为png格式，命名规则，word名和图片名必须一致
        Args:
            src_dir (str, optional): 源文档目录. Defaults to "docs".
            save_dir (str, optional): 存储文档目录. Defaults to "results".

        Returns:
            _type_: _description_
        """
        self.src_dir = src_dir
        self.save_dir = save_dir
        self.image_dir = image_dir
        self.img_flag = img_flag

        self.l_words = self.get_all_docs() # 获取所有word文档路径
        self.d_data = self.get_images()
    def get_images(self):
        data = {}
        for d in self.l_words:
            if d in data:
                raise "重名文件，请检查文件！"
            image_name = d.replace(self.src_dir, self.image_dir).replace('.docx', '.png')
            print(f"image_name==>{image_name}")
            if not os.path.exists(image_name):
                raise FileNotFoundError(errno.ENOENT, os.strerror(errno.ENOENT), image_name)
            data[d] = image_name
        return data
        
    def get_all_docs(self):
        print(self.src_dir)
        
        docs = glob(f'{self.src_dir}/*.docx')
        print(docs)
        if len(docs) > 0:
            return docs
        return None

    def process_docs(self):
        for d, img in self.d_data.items():
            print("process==>")
            print(d, img)
            docx = Document(d) #docx文件的地址
            tables = docx.tables # 获取所有表格
            for t in tables:
                print(dir(t))
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            print(f"p=>{p.text}")
                            text = p.text
                            if self.has_img_flag(text):
                                # 插入图片
                                p.text = p.text.replace('<<img1>>', '')
                                run = p.add_run('')
                                run.add_break()
                                print("=====>"*10)
                                pic = run.add_picture(img, width=Inches(1.2))
            save_path = d.replace(self.src_dir, self.save_dir)
            docx.save(save_path)

    def has_img_flag(self, text):
        if self.img_flag in text:
            return True
        return False


prod = SaveImgInTable()
prod.process_docs()
print(prod.d_data)

print("处理完成！")